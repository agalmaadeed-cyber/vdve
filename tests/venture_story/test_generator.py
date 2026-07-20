import io

from docx import Document

from venture_story.docx_renderer import BOLD_END, BOLD_START
from venture_story.generator import generate_venture_story


def _make_field(field_code, value, evidence_label="ESTIMATE"):
    return {"field_code": field_code, "value": value, "evidence_label": evidence_label}


def _build_dossier():
    codes = (
        "A1", "A2", "A3", "A4", "A5",
        "C1", "C2", "C3", "C4", "C5",
        "B1", "B2", "B3", "B4", "B5", "B6", "B7",
        "D1", "D2", "D3", "D4", "D5", "D6",
        "F1", "F2", "F3", "F4",
    )
    fields = {code: _make_field(code, f"Value for {code}.") for code in codes}
    return {
        "dossier_id": "DS-TEST",
        "version": 2,
        "source": {
            "uh_idea_name": "TestVenture",
            "uh_sector": "Productivity SaaS",
            "uh_final_score": "43/50",
            "uh_final_decision": "Test with reservations.",
            "uh_next_step": "Run a manual pilot first.",
        },
        "sections": {
            "opportunity": {"a1": fields["A1"], "a2": fields["A2"], "a3": fields["A3"], "a4": fields["A4"], "a5": fields["A5"]},
            "solution": {"c1": fields["C1"], "c2": fields["C2"], "c3": fields["C3"], "c4": fields["C4"], "c5": fields["C5"]},
            "customer_market": {
                "b1": fields["B1"], "b2": fields["B2"], "b3": fields["B3"], "b4": fields["B4"],
                "b5": fields["B5"], "b6": fields["B6"], "b7": fields["B7"],
            },
            "business_model": {"d1": fields["D1"], "d2": fields["D2"], "d3": fields["D3"], "d4": fields["D4"], "d5": fields["D5"], "d6": fields["D6"]},
            "success_definition": {"f1": fields["F1"], "f2": fields["F2"], "f3": fields["F3"], "f4": fields["F4"]},
        },
    }


def _build_cycle_record():
    return {
        "cycle_record_id": "cr-test-1",
        "dossier_id": "DS-TEST",
        "dossier_version": 2,
        "claim_field_codes": ["A1", "B1"],
        "unknown_field_codes": [],
        "approved_parameters": {},
        "stress_test_results": [
            {
                "test_id": "ST-01", "test_type": "quantitative_shock", "category": "cost",
                "source": "fixed_library", "status": "COMPLETED", "outcome": "SURVIVES",
                "target_hypothesis_id": None, "rationale": None, "severity": None,
            },
            {
                "test_id": "GEN-A1", "test_type": "qualitative_probe", "category": "demand",
                "source": "generated", "status": "COMPLETED", "outcome": None,
                "target_hypothesis_id": "A1", "rationale": "Plausible but unverified.", "severity": "MEDIUM",
            },
        ],
        "kill_status": "No concern",
        "ceiling_result": {"ceiling": "Pass with Conditions", "triggered_by": ["Pass with Conditions:stress_test_breaks:ST-04"]},
        "recommendation": {
            "outcome": "Pass with Conditions",
            "status": "LLM_RECOMMENDED",
            "narrative": "This is the recommendation narrative.",
            "payload": {"conditions": [{"hypothesis_id": "A1", "condition": "Validate demand via field interviews."}]},
        },
    }


def _build_gate4_verdict():
    return {
        "result": "PASS",
        "checks": [
            {"criterion": 1, "id": "outcome_in_advance_range", "description": "Outcome in advance range", "applicable": True, "passed": True, "evidence": {}},
            {"criterion": 2, "id": "version_current", "description": "Version is current", "applicable": True, "passed": True, "evidence": {}},
        ],
        "reason_codes": [],
        "block_routes_to": None,
        "founder_signoff": None,
        "checked_at": "2026-07-20T00:00:00+00:00",
        "chain_fingerprint": {},
    }


def _build_scenarios():
    def _metrics(price):
        return {
            "price_per_unit": price, "variable_cost_per_unit": 0.02, "CAC": 20.0,
            "avg_customer_lifetime_months": 9.0, "monthly_burn": 30.0, "budget": 1000.0,
            "gross_margin": 0.999, "LTV": 215.82, "LTV_to_CAC": 10.79,
            "payback_period": 0.41, "runway_months": 33.33, "breakeven_customers": 0.61,
        }
    return {
        "conservative": _metrics(19.2),
        "base": _metrics(24.0),
        "optimistic": _metrics(28.8),
    }


SIGNED_OFF_AT = "2026-07-20T12:00:00+00:00"


def test_generate_venture_story_full_pipeline_smoke_test():
    dossier = _build_dossier()
    cycle_record = _build_cycle_record()
    gate4_verdict = _build_gate4_verdict()
    scenarios = _build_scenarios()

    docx_bytes = generate_venture_story(dossier, cycle_record, gate4_verdict, SIGNED_OFF_AT, scenarios)

    assert isinstance(docx_bytes, bytes)
    assert len(docx_bytes) > 0

    doc = Document(io.BytesIO(docx_bytes))  # opens as a valid .docx

    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"]
    for i in range(1, 12):
        assert any(f"{i}." in h for h in heading_texts), f"no heading found for section {i}"

    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text

    assert BOLD_START not in full_text
    assert BOLD_END not in full_text
