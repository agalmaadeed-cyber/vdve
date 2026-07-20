import io

from docx import Document

from venture_story.docx_renderer import (
    BOLD_END,
    BOLD_START,
    _iter_bold_spans,
    render_markdown_to_docx,
    split_bidi_runs,
)

# The real A1 sentence from fixtures/DS-0FE02838.json -- the exact string the
# 2026-07-20 isolated spike ran against.
A1_ARABIC = (
    "المدير يعيش في دوامة تنظيمية متكررة كل 3-8 أسابيع: فوضى → أزمة → محاولة "
    "تنظيم → إرهاق → هجر → فوضى من جديد. الحلول الموجودة تضع عبء البناء والصيانة "
    "على المستخدم نفسه."
)

IDEA_NAME_MIXED = "ChiefBot — المساعد التنفيذي على واتساب/تيليجرام"


# --- split_bidi_runs() ---

def test_pure_english_text_is_one_non_arabic_segment():
    text = "This is plain English text."
    result = split_bidi_runs(text)
    assert result == [(text, False)]


def test_pure_arabic_text_is_one_arabic_segment_full_text_preserved():
    result = split_bidi_runs(A1_ARABIC)
    assert len(result) == 1
    segment, is_arabic = result[0]
    assert is_arabic is True
    assert segment == A1_ARABIC


def test_english_shell_around_arabic_quote_produces_three_ordered_segments():
    text = (
        'Section 3 -- Opportunity Definition. Problem (verbatim from field A1): "'
        + A1_ARABIC
        + '" This is the founder-declared problem statement.'
    )
    result = split_bidi_runs(text)
    assert [is_arabic for _, is_arabic in result] == [False, True, False]
    assert "".join(segment for segment, _ in result) == text


def test_mixed_idea_name_produces_two_ordered_segments():
    result = split_bidi_runs(IDEA_NAME_MIXED)
    assert [is_arabic for _, is_arabic in result] == [False, True]
    assert "".join(segment for segment, _ in result) == IDEA_NAME_MIXED


# --- _iter_bold_spans() ---

def test_bold_wrapped_span_yields_text_true_with_sentinel_stripped():
    text = f"{BOLD_START}The Problem:{BOLD_END} some plain text"
    result = list(_iter_bold_spans(text))
    assert result[0] == ("The Problem:", True)
    assert result[1] == (" some plain text", False)


def test_literal_double_asterisk_passes_through_completely_unchanged():
    # Simulates a founder's own raw text containing literal "**", which must
    # never be treated as this generator's own bold marker (point (g) fix).
    text = 'A quote that says "**this is bold in markdown**" verbatim.'
    result = list(_iter_bold_spans(text))
    assert result == [(text, False)]


# --- render_markdown_to_docx() round-trip ---

def _rtl_runs(paragraph):
    rtl_flags = []
    for run in paragraph.runs:
        rPr = run._element.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
        )
        has_rtl = rPr is not None and rPr.find(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rtl"
        ) is not None
        rtl_flags.append(has_rtl)
    return rtl_flags


def test_round_trip_mixed_bidi_paragraph_has_three_runs_rtl_on_middle_only():
    mixed_sentence = (
        'Section 3 -- Opportunity Definition. Problem (verbatim from field A1): "'
        + A1_ARABIC
        + '" This is the founder-declared problem statement.'
    )
    markdown = f"# Title\n\n{mixed_sentence}\n"

    docx_bytes = render_markdown_to_docx(markdown)
    doc = Document(io.BytesIO(docx_bytes))

    body_paragraphs = [p for p in doc.paragraphs if A1_ARABIC in p.text]
    assert len(body_paragraphs) == 1
    paragraph = body_paragraphs[0]

    assert len(paragraph.runs) == 3
    assert _rtl_runs(paragraph) == [False, True, False]


def test_round_trip_bold_label_is_bold_literal_asterisks_are_not():
    markdown = (
        "# Title\n\n"
        f"{BOLD_START}Label:{BOLD_END} a quote containing literal **not our marker** text.\n"
    )
    docx_bytes = render_markdown_to_docx(markdown)
    doc = Document(io.BytesIO(docx_bytes))

    body_paragraphs = [p for p in doc.paragraphs if "Label:" in p.text]
    assert len(body_paragraphs) == 1
    paragraph = body_paragraphs[0]

    bold_run = next(r for r in paragraph.runs if r.text == "Label:")
    assert bold_run.bold is True

    literal_run = next(r for r in paragraph.runs if "**not our marker**" in r.text)
    assert literal_run.bold is not True


def test_no_bold_sentinel_leaks_into_any_paragraph_or_table_cell_text():
    markdown = (
        "# Title\n\n"
        f"{BOLD_START}Label:{BOLD_END} some text.\n\n"
        "| Metric | Value |\n"
        "|---|---|\n"
        f"| {BOLD_START}Row{BOLD_END} | 1 |\n"
    )
    docx_bytes = render_markdown_to_docx(markdown)
    doc = Document(io.BytesIO(docx_bytes))

    for paragraph in doc.paragraphs:
        assert BOLD_START not in paragraph.text
        assert BOLD_END not in paragraph.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                assert BOLD_START not in cell.text
                assert BOLD_END not in cell.text


def test_round_trip_opens_without_error_and_pipe_table_has_right_shape_and_bold_header():
    markdown = (
        "# Title\n\n"
        "## Section\n\n"
        "| Metric | Conservative | Base | Optimistic |\n"
        "|---|---|---|---|\n"
        "| Price | 19.20 | 24.00 | 28.80 |\n"
        "| CAC | 24.00 | 20.00 | 16.00 |\n"
    )
    docx_bytes = render_markdown_to_docx(markdown)
    doc = Document(io.BytesIO(docx_bytes))  # (i) opens without error

    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 3  # header + 2 data rows
    assert len(table.columns) == 4

    header_row = table.rows[0]
    for cell in header_row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                assert run.bold is True
