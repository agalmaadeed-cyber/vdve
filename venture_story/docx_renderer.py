"""
The one controlled Markdown -> .docx conversion step (Decision
P1.0.10 point 2). Parses the small, fixed Markdown subset produced by
markdown_builder.py (headings, paragraphs, pipe tables) and is the
ONLY place in this package that deals with bidi/RTL rendering.

Bidi approach, verified before any of this was built (2026-07-20
isolated spike, see phase1_decisions_log.md P1.0.10 point 3 and this
packet's own S0(b)): a run-level w:rtl + w:cs marking on any
contiguous Arabic-script segment inside an otherwise-LTR paragraph
renders correctly -- confirmed visually via python-docx + LibreOffice
headless render against the real DS-0FE02838 A1 sentence embedded in
an English-shell paragraph.

split_bidi_runs() is the generalized form of that exact verified
case: it tokenizes text into words/whitespace, tracks a running
"current script mode" (Arabic vs. not), flips to Arabic on any token
containing an Arabic-range character, flips back to non-Arabic only
on a token containing a Latin LETTER (digits, punctuation, arrows,
and whitespace are treated as neutral and stay attached to whichever
run is already open -- this is what keeps "3-8" and the arrows inside
the real A1 sentence part of the same Arabic run, exactly matching
what the spike's variant B/D already proved renders correctly as one
marked run, rather than fragmenting into many tiny runs at every
digit or space).
"""

from __future__ import annotations

import io
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Explicit \uXXXX escapes, not literal Arabic characters, so this
# regex survives copy/paste and file transfer byte-for-byte -- same
# "encoding is a known failure point in this project" caution as
# P1.0.7's double-decode note.
_ARABIC_CHAR = re.compile(
    "[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF\uFB50-\uFDFF\uFE70-\uFEFF]"
)
_LATIN_LETTER = re.compile(r"[A-Za-z]")
_TOKEN = re.compile(r"\S+|\s+")

# Bold-label sentinels used ONLY by this project's own markdown_builder.py
# template labels -- deliberately NOT the literal "**" GFM bold marker.
# Real Dossier raw_dossier_text values were found during pre-flight testing
# to themselves sometimes contain literal "**" (founders/LLM-assisted
# drafting tools produced markdown-styled free text) -- reusing "**" as our
# own bold marker would either wrongly bold-format a founder's verbatim
# quote or require escaping logic that risks altering quoted text, both of
# which violate this generator's "never reworded" quoting rule (S0(c)).
# Private-Use-Area characters can never appear in real Dossier text, so
# there is no collision risk, and no escaping is ever needed.
BOLD_START = "\ue000"
BOLD_END = "\ue001"


def _iter_bold_spans(text: str):
    """Yields (text, is_bold) chunks, splitting on the BOLD_START/BOLD_END
    sentinel pair. Any BOLD_START with no matching BOLD_END is treated as
    plain text (never crashes, never silently drops content)."""
    parts = re.split(f"({BOLD_START}.*?{BOLD_END})", text)
    for part in parts:
        if not part:
            continue
        if part.startswith(BOLD_START) and part.endswith(BOLD_END):
            yield part[len(BOLD_START):-len(BOLD_END)], True
        else:
            yield part, False


def split_bidi_runs(text: str) -> list[tuple[str, bool]]:
    """
    Splits text into (segment, is_arabic) tuples, merging adjacent
    tokens that share the same script mode. See module docstring for
    the exact flip rule. Never called on an empty string by this
    module's own callers, but returns a single non-Arabic segment for
    one if it ever is -- no crash on empty content.
    """
    if not text:
        return [(text, False)]
    segments: list[list] = []
    mode = False
    for tok in _TOKEN.findall(text):
        if _ARABIC_CHAR.search(tok):
            mode = True
        elif _LATIN_LETTER.search(tok):
            mode = False
        # else: neutral token (digits/punctuation/whitespace/arrows) --
        # carries forward the current mode unchanged.
        if segments and segments[-1][1] == mode:
            segments[-1][0] += tok
        else:
            segments.append([tok, mode])
    return [(t, a) for t, a in segments]


def _mark_run_rtl(run) -> None:
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "true")
    rPr.append(rtl)
    rPr.append(OxmlElement("w:cs"))


def _add_bidi_runs(paragraph, text: str) -> None:
    for span_text, is_bold in _iter_bold_spans(text):
        for segment, is_arabic in split_bidi_runs(span_text):
            if not segment:
                continue
            run = paragraph.add_run(segment)
            if is_arabic:
                _mark_run_rtl(run)
            if is_bold:
                run.bold = True


def _parse_table(rows: list[str]) -> list[list[str]]:
    parsed: list[list[str]] = []
    for r in rows:
        cells = [c.strip() for c in r.strip().strip("|").split("|")]
        if cells and all(set(c) <= {"-", ":", ""} for c in cells):
            continue  # markdown separator row, e.g. |---|---|
        parsed.append(cells)
    return parsed


def _parse_blocks(markdown_text: str) -> list[dict]:
    blocks: list[dict] = []
    table_rows: list[str] = []

    def _flush_table() -> None:
        if table_rows:
            blocks.append({"type": "table", "rows": _parse_table(table_rows)})
            table_rows.clear()

    for raw_line in markdown_text.split("\n"):
        line = raw_line.rstrip()
        if line.startswith("|"):
            table_rows.append(line)
            continue
        _flush_table()
        if not line.strip():
            continue
        if line.startswith("# "):
            blocks.append({"type": "h1", "text": line[2:].strip()})
        elif line.startswith("## "):
            blocks.append({"type": "h2", "text": line[3:].strip()})
        elif line.startswith("### "):
            blocks.append({"type": "h3", "text": line[4:].strip()})
        else:
            blocks.append({"type": "paragraph", "text": line})
    _flush_table()
    return blocks


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=0, cols=n_cols)
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass  # falls back to the default table look -- cosmetic only, never blocks generation
    for i, row_cells in enumerate(rows):
        row = table.add_row()
        for j, cell_text in enumerate(row_cells[:n_cols]):
            cell_paragraph = row.cells[j].paragraphs[0]
            for span_text, is_bold in _iter_bold_spans(cell_text):
                for segment, is_arabic in split_bidi_runs(span_text):
                    if not segment:
                        continue
                    run = cell_paragraph.add_run(segment)
                    if is_arabic:
                        _mark_run_rtl(run)
                    if is_bold or i == 0:
                        run.bold = True


def render_markdown_to_docx(markdown_text: str) -> bytes:
    """
    Pure function: Markdown string in, .docx bytes out. Writes
    nothing to disk itself -- the caller (generator.py, and
    ultimately Streamlit's download_button) owns what happens to the
    bytes.
    """
    doc = Document()
    for block in _parse_blocks(markdown_text):
        if block["type"] == "h1":
            doc.add_heading(block["text"], level=0)
        elif block["type"] == "h2":
            doc.add_heading(block["text"], level=1)
        elif block["type"] == "h3":
            doc.add_heading(block["text"], level=2)
        elif block["type"] == "table":
            _add_table(doc, block["rows"])
        elif block["type"] == "paragraph":
            paragraph = doc.add_paragraph()
            _add_bidi_runs(paragraph, block["text"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
